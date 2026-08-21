from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


BODY_EAST_ASIA = "微软雅黑"
HEADING_EAST_ASIA = "楷体"
LATIN_FONT = "Times New Roman"
MONO_FONT = "Consolas"
INK = RGBColor(31, 55, 78)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)


def set_run_fonts(run, east_asia: str, latin: str = LATIN_FONT) -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def replace_paragraph(document: Document, old: str, new: str) -> Paragraph:
    paragraph = next((item for item in document.paragraphs if item.text == old), None)
    if paragraph is None:
        raise ValueError(f"Paragraph not found: {old[:80]}")
    paragraph.text = new
    return paragraph


def paragraph_by_text(document: Document, text: str) -> Paragraph:
    paragraph = next((item for item in document.paragraphs if item.text == text), None)
    if paragraph is None:
        raise ValueError(f"Paragraph not found: {text[:80]}")
    return paragraph


def remove_paragraph(document: Document, text: str) -> None:
    paragraph = paragraph_by_text(document, text)
    paragraph._element.getparent().remove(paragraph._element)


def remove_paragraph(document: Document, text: str) -> None:
    paragraph = paragraph_by_text(document, text)
    paragraph._element.getparent().remove(paragraph._element)


def insert_after(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_sequence(anchor: Paragraph, items: list[tuple[str, str]]) -> Paragraph:
    current = anchor
    for style, text in items:
        current = insert_after(current, text, style)
    return current


def insert_table_after(
    document: Document,
    anchor: Paragraph,
    rows: list[list[str]],
    widths: list[float],
) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    for column, width in zip(table.columns, widths):
        column.width = Inches(width)
    for row_index, values in enumerate(rows):
        row = table.rows[0] if row_index == 0 else table.add_row()
        for column_index, value in enumerate(values):
            row.cells[column_index].text = value
            row.cells[column_index].width = Inches(widths[column_index])
    anchor._p.addnext(table._tbl)


def set_table_data(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row, values in zip(table.rows, rows):
        if len(row.cells) != len(values):
            raise ValueError("Table column count does not match replacement data")
        for cell, value in zip(row.cells, values):
            cell.text = value


def table_by_headers(document: Document, headers: list[str]):
    for table in document.tables:
        if not table.rows or len(table.rows[0].cells) != len(headers):
            continue
        values = [cell.text.strip() for cell in table.rows[0].cells]
        if values == headers:
            return table
    raise ValueError(f"Table not found: {headers}")


def add_tool_section(document: Document) -> None:
    anchor = paragraph_by_text(
        document,
        "系统跨越浏览器、应用服务器、外部学术服务和高性能计算集群四个信任域。浏览器端不保存 LLM 密钥、SSH 私钥和 POTCAR；FastAPI 对外部文本、结构文件和审查决定实施白名单校验；学术接口返回的元数据仅作为待核验输入；集群操作由服务器端受控模板产生，并分别受只读预检、远程写入和作业提交开关控制。网页真实远程操作默认关闭，只有管理员显式开启总开关和底层开关后才可进入上传与提交阶段，且无法绕过独立人工确认。",
    )
    heading = insert_after(anchor, "2.4 核心工具与职责", "Heading 2")
    lead = insert_after(
        heading,
        "系统不要求单一模型承担全部科研任务，而是根据工具的可验证能力分工：Kimi 处理语义与解释，学术接口处理来源，ASE/pymatgen 与 CGCNN 处理结构和低成本预测，VASP 与 Slurm 处理高成本计算，LangGraph、FastAPI 和持久化记录负责把这些工具组织为可恢复流程。",
    )
    rows = [
        ["工具类别", "主要工具", "在 Catalyst Agent 中的职责"],
        ["语义模型", "Kimi K3", "识别自然语言任务、形成检索表达、解释术语、辅助 VASP 参数咨询和生成任务报告；不直接获得远程执行权限。"],
        ["学术证据", "本地高熵数据库、Crossref、Semantic Scholar", "本地优先召回、在线补充、正式元数据核验、摘要与开放链接补充。"],
        ["工作流编排", "LangGraph、SQLite Checkpoint", "节点路由、人工中断、状态持久化和统一恢复。"],
        ["服务与界面", "FastAPI、assistant-ui、Next.js", "任务接口、聊天台、时间线、审查卡、历史任务、文件和 DFT 监控。"],
        ["结构建模", "ASE、pymatgen", "POSCAR/CIF 解析，FCC bulk、(111) slab、吸附位点和结构派生。"],
        ["机器学习", "CGCNN", "形成能预筛、临时训练与生产模型结果比较。"],
        ["第一性原理", "VASP、PAW/POTCAR", "bulk、clean slab 和吸附体系弛豫及能量输出。"],
        ["高性能计算", "OpenSSH、Slurm", "只读预检、上传、SHA-256 校验、sbatch、状态查询、下载和恢复。"],
        ["运行配置", "Docker Compose", "定义前后端服务、数据卷与 HPC 只读挂载；敏感配置与许可数据不写入镜像。"],
    ]
    insert_table_after(document, lead, rows, [1.05, 1.55, 3.9])


def add_workflow_completeness_section(document: Document) -> None:
    anchor = paragraph_by_text(
        document,
        "C 阶段把候选设计、结构生成和计算验证组织为递进的资源分配过程。C1-C4 完成候选约束、评分、五元组合生成与人工选择；C5 构造 FCC bulk；C6 在生产 CGCNN、临时训练模型和 Bulk DFT 路径之间形成可审查的形成能来源；C7 联合形成能、原子尺寸差 δ 和 Ω 参数筛选可升级结构；C8-C10 生成并审查（111）slab 与 VASP 五文件；C11 执行预检、上传、SHA-256 校验、Slurm 提交、监控、下载和结果解析；C12.1-C12.7 从弛豫后的 clean slab 选择单一中间体与位点，建立吸附结构、准备并执行吸附 DFT，最后核对三项能量及参考来源。每次阶段升级均提高计算成本和结论具体性，因此需要更严格的身份校验与人工确认。",
    )
    heading = insert_after(anchor, "3.5 全流程节点链与阶段交付", "Heading 2")
    lead = insert_after(
        heading,
        "三条入口在不同位置进入同一科学对象链，但所有进入 DFT 的对象最终都遵循结构身份、输入审查、远程确认、结果解析和人工接受规则。完整节点链及其可见交付如下。",
    )
    rows = [
        ["阶段", "核心职责", "用户可见交付"],
        ["A1-A4", "意图解析、能力判断、三入口路由、阶段计划", "任务理解摘要、路线与下一阶段"],
        ["B1-B6", "本地优先召回、在线补充、评分、断言提取与人工审查", "文献评分、双语题名/摘要、证据与审查记录"],
        ["C1-C4", "约束、六维评分、五元组合生成与选择", "最多三个候选、雷达图和人工选择"],
        ["C5-C7", "FCC bulk、形成能来源与固溶体稳定性预筛", "结构预览、eV/atom、δ、Ω 与升级决定"],
        ["C8-C10", "切取 (111) slab、几何质检、VASP 五文件审查", "slab 三维结构、参数预览与可控修改"],
        ["C11", "集群预检、上传、摘要校验、提交、监控、下载与解析", "连接状态、Slurm 状态、日志、结果文件"],
        ["C12.1-C12.7", "中间体和位点选择、吸附结构、吸附 DFT、三能量计算", "吸附结构、位点、Eads 公式、来源与人工结论"],
    ]
    insert_table_after(document, lead, rows, [1.15, 3.15, 2.2])
    paragraph_by_text(document, "3.5 节点设计规范").text = "3.6 节点设计规范"


def apply_content_updates(document: Document) -> None:
    replacements = {
        "就当前实现而言，代码库已经形成包含45个节点的主工作流，并在文献采用、候选选择、执行范围、结构升级、slab审核、VASP输入、远程上传和Slurm提交等环节设置显式中断。FastAPI已提供任务创建、状态查询和低风险人工审查接口，assistant-ui前端已接入文献审查、候选选择、C阶段执行范围和C7结构升级等人工门。由于API进程当前强制关闭集群预检、远程写入和作业提交，网页端仍属于受限运行阶段。因此，本文将Catalyst Agent界定为可恢复的计算科研工作流与决策支持环境，而不是无人监督的自主实验室。":
            "截至 2026 年 7 月，代码库已经形成覆盖 A1-A4、B1-B6、C1-C12.7 的主工作流，并在文献、候选、执行范围、结构升级、slab、VASP 输入、远程上传、Slurm 提交和吸附能等环节设置显式人工中断。FastAPI 已提供任务、审查、历史恢复、文件、结构、报告、训练和作业状态接口；assistant-ui/Next.js 工作台已呈现聊天记录、完整时间线、阶段摘要、统一审查卡、任务历史、文件查看、三维结构和 DFT 监控。标题栏通过最小 Kimi 请求和只读 SSH 探针显示真实连接状态。网页远程操作采用显式总开关，并继续强制执行 UPLOAD <task_id> 与 SUBMIT <task_id> 两道确认门。因此，Catalyst Agent 是可恢复、可审查的计算科研工作流和决策支持环境，而不是无人监督的自主实验室。",
        "本研究面向从事电化学、催化、材料计算和高熵合金研究的科研人员，主要服务于以下连续场景：从研究问题出发建立可核验的文献证据；在明确的元素约束下生成候选组成并进行软评分；针对五元组成或外部POSCAR/CIF结构构造bulk、slab及吸附模型；最后跟踪计算作业、检查收敛性并形成可追溯的结果记录。该部分描述的是研究流程中的适用范围，而非面向普通用户的产品功能清单。":
            "系统面向电化学、催化、高熵合金和材料计算研究者，也允许缺少 VASP、Linux 或超算经验的使用者从自然语言和结构文件开始。领域熟悉者可以直接指定元素组合、POSCAR/CIF、执行范围和参数修订；计算经验较少的用户则通过术语解释、推荐选项、阶段摘要和人工确认理解每一步的目的与风险。系统降低的是工具编排和重复操作门槛，并不降低形成科研结论所需的证据标准。",
        "系统跨越浏览器、应用服务器、外部学术服务和高性能计算集群四个信任域。浏览器端不保存 LLM 密钥、SSH 私钥和 POTCAR；FastAPI 对外部文本、结构文件和审查决定实施白名单校验；学术接口返回的元数据仅作为待核验输入；集群操作由服务器端受控模板产生，并分别受只读预检、远程写入和作业提交开关控制。当前网页 API 在进程启动时强制关闭三类集群能力，因而即使环境变量配置错误，也不会由普通网页任务产生远程副作用。":
            "系统跨越浏览器、应用服务器、外部学术服务和高性能计算集群四个信任域。浏览器端不保存 LLM 密钥、SSH 私钥和 POTCAR；FastAPI 对外部文本、结构文件和审查决定实施白名单校验；学术接口返回的元数据仅作为待核验输入；集群操作由服务器端受控模板产生，并分别受只读预检、远程写入和作业提交开关控制。网页真实远程操作默认关闭，只有管理员显式开启总开关和底层开关后才可进入上传与提交阶段，且无法绕过独立人工确认。",
        "主工作流以研究问题的语义状态而非固定页面顺序驱动执行。task_analysis 首先确定反应体系、材料家族、用户约束和计算意图，随后在 normal、direct_c 与 external_c 三条入口之间选择。normal 路径依次建立文献证据与候选体系；direct_c 将用户明确给出的五元组成视为固定化学体系，仅生成不同原子排布；external_c 对实际 POSCAR/CIF 进行解析，并依据是否具有形成能决定进入预测、Bulk DFT 或稳定性阶段。三条路径在结构与计算对象上重新汇合，但保留各自的证据来源和 scientific_scope。":
            "主工作流以研究问题的语义状态而非固定页面顺序驱动执行，并支持三条入口。normal 路径从自然语言任务出发，完整执行 A、B、C 阶段；direct_c 在用户明确提出五元高熵合金建模时跳过候选文献链，直接建立该组成的结构对象；external_c 接收用户给出的 POSCAR/CIF 路径，先由 CGCNN 预测形成能并执行稳定性判据，再由用户决定是否升级 slab 与 DFT。三条路径在结构、VASP、超算和结果审查环节汇合，同时保留各自的证据来源、scientific_scope 和 task_id。",
        "A 阶段承担自然语言与科研状态之间的规范化转换。任务分析输出反应类型、目标产物、材料家族、外部结构请求及所需能力；反应档案用于约束当前实现能够讨论的中间体和计算范围；能力门判断请求是否超出已有工具；路由器决定入口类型；规划器选择是否需要文献检索以及后续 C 阶段深度。语言模型失败时，规则回退保留可运行性，但必须记录 analysis_mode 与警告，不能把回退结果表述为模型已完成复杂科学推理。":
            "A 阶段承担自然语言与科研状态之间的规范化转换。A1 提取反应、目标产物、材料家族、元素、外部结构和计算意图；A2 检查文献、建模、预测、DFT 与 OER 等反应能力；A3 在 normal、direct_c 和 external_c 之间路由；A4 生成阶段计划、证据需求和执行深度。Kimi 用于处理宽泛表达和形成更精准的检索输入，模型失败时由确定性规则回退并记录 analysis_mode 与警告，不能把规则结果表述为模型已完成复杂科学推理。",
        "B 阶段围绕“论文记录—原文断言—人工采用”建立证据链。本地 SQLite/JSON 语料首先执行结构化过滤和词法召回，在线策略再根据独立 DOI、题名和高质量结果数量决定是否补充 Crossref 与 Semantic Scholar。在线记录按 DOI 优先、规范化题名次之的规则合并，并保存单源或交叉核验等级。抽取服务将材料组成、反应、性能条件和结论拆分为可审查断言；只有研究者接受的论文及断言进入 RAG 摘要和候选生成。若证据全部被拒绝，系统在限定轮次内改变查询并保留拒绝历史，避免用重复结果制造表面上的证据充足。":
            "B 阶段按 B1-B6 建立“论文记录—原文断言—人工采用”证据链。本地高熵合金数据库承担主要召回份额，B4 在线检索通过 Crossref 与 Semantic Scholar 补充并核验元数据，避免在线接口限流拖慢全部流程。评分同时考虑元数据、任务相关性、组成与高熵证据、可引用断言、期刊影响和常用过渡金属覆盖；前端输出各维度得分、六维雷达图、英文原文与中文翻译。只有人工接受的论文和科学断言进入候选生成；若全部拒绝，系统改变查询并排除上一轮结果。单源可追溯记录允许继续，但不得伪装成双源核验。",
        "C 阶段把候选设计、结构生成和计算验证组织为递进的资源分配过程。C1-C4 在五元、FCC 32 原子及元素排除规则下生成并排序候选；C5 构造 bulk 结构；C6 根据 CGCNN 元素域选择预测或 Bulk DFT；C7 联合形成能、原子尺寸差 δ 和 Ω 参数决定是否具备 slab 建模资格；C8-C10 生成并审查（111）slab 及 VASP 五文件；C11 通过预检、上传、提交、监控、下载和解析管理长周期作业；C12 从弛豫后的 clean slab 出发生成单一中间体的位点结构，并以三项能量差计算吸附能。每次阶段升级均提高计算成本和结论具体性，因此需要更严格的身份校验与人工确认。":
            "C 阶段把候选设计、结构生成和计算验证组织为递进的资源分配过程。C1-C4 完成候选约束、评分、五元组合生成与人工选择；C5 构造 FCC bulk；C6 在生产 CGCNN、临时训练模型和 Bulk DFT 路径之间形成可审查的形成能来源；C7 联合形成能、原子尺寸差 δ 和 Ω 参数筛选可升级结构；C8-C10 生成并审查（111）slab 与 VASP 五文件；C11 执行预检、上传、SHA-256 校验、Slurm 提交、监控、下载和结果解析；C12.1-C12.7 从弛豫后的 clean slab 选择单一中间体与位点，建立吸附结构、准备并执行吸附 DFT，最后核对三项能量及参考来源。每次阶段升级均提高计算成本和结论具体性，因此需要更严格的身份校验与人工确认。",
        "Human-in-the-loop 在本系统中不是界面层附加功能，而是科研状态转移的组成部分。LangGraph interrupt() 将当前审查对象、允许动作和恢复位置固化到 Checkpoint；FastAPI 提交审查时校验 task_id、review_id、review_type、对象集合和幂等键，再以 Command(resume=decision) 恢复同一 thread_id。当前网页端仅开放文献、候选、执行范围和 C7 结构升级等低风险审查，高风险远程操作仍由服务端拒绝，从而使交互能力的开放程度与其潜在副作用相匹配。":
            "Human-in-the-loop 在本系统中不是界面层附加功能，而是科研状态转移的组成部分。LangGraph interrupt() 将审查对象、允许动作和恢复位置固化到 Checkpoint；FastAPI 校验 task_id、review_id、review_type、对象集合和幂等键，再恢复同一 thread_id。前端已经统一呈现 B6、C4、C7、C9、C10、C11 和 C12.7 等审查卡。远程操作只有在管理员开启 WEB_REMOTE_OPERATIONS_ENABLED 及底层集群开关后才可进入，并继续要求精确的 UPLOAD 与 SUBMIT 短语，因此界面便利性不会绕过外部副作用边界。",
        "接口设计应明确区分当前已实现能力和后续规划能力。当前接口包括任务创建、任务状态查询和低风险人工审查；结构上传、Slurm作业管理及科研产物下载属于扩展接口，投入使用前需要补充身份认证、任务归属、文件隔离、授权下载和高风险确认机制。":
            "当前 FastAPI 提供任务创建、状态查询、人工审查、历史恢复、任务归档、文件列表与预览、结构查看、CGCNN 训练、Slurm 状态与日志、报告生成及 Kimi/HPC 连接检查。前端消费这些接口，形成聊天台、阶段时间线、审查卡与右侧任务工作台。",
        "前端页面围绕“任务建立—证据审查—候选选择—结构与计算状态—结果回溯”组织。assistant-ui/Next.js负责呈现研究问题、证据摘要、候选信息、任务进度和人工审查对象，并将研究者决定提交给FastAPI。当前已接入文献审查、候选选择、执行范围和C7结构升级等人工门；POSCAR/CIF上传、作业中心、结果下载和三维结构可视化属于后续建设内容。页面只负责展示和收集决定，不改变后端确定的科学判据和安全开关。":
            "前端采用桌面端三栏组织：左侧为 A1-C12.7 时间线，中部为自然语言对话、逐节点摘要和人工审查卡，右侧为任务历史、恢复、文件、结构、报告和 DFT 监控。候选可显示评分雷达图，bulk、slab 与吸附结构提供 ASE 风格三维查看，VASP 文件可按安全范围预览和受控修订，Kimi/HPC 状态显示在标题栏。工作流内容持续保留在聊天台中，页面只展示和收集决定，不修改后端科学判据或安全开关。",
        "8 容器化与部署方案（待工程实现）": "8 当前运行与容器配置",
        "项目当前采用 Python 虚拟环境与 Next.js 双终端方式运行，代码库中尚未形成正式 Dockerfile、docker-compose.yml 或 Compose 生产配置。因此，本章属于基于现有依赖和信任边界提出的部署设计，而非对已经完成容器化的描述。建议将前端、API、作业监控和反向代理分离，以避免网页进程直接持有集群权限，并为不同服务配置独立健康检查、资源限制和文件系统权限。":
            "代码库提供 Dockerfile.backend、frontend/Dockerfile、docker-compose.yml、docker-compose.hpc.yml 和 .dockerignore。默认 Compose 关闭真实超算操作，用于安全演示；HPC 叠加配置以只读方式挂载 SSH 私钥、known_hosts 与许可 PBE 数据，并继续复用网页人工确认门。",
        "建议的生产拓扑为“浏览器—HTTPS 反向代理—Next.js—FastAPI/LangGraph—受控持久化—HPC”。反向代理承担 TLS、认证、限流和请求大小控制；Next.js 通过内部网络调用 FastAPI；API 服务不直接暴露公网；监控服务仅访问持久化作业记录和最小化 SSH 凭据；POTCAR 以只读许可卷挂载，计算数据以独立持久卷保存。若前后端部署在不同节点，还需建立可信服务间认证，而不能仅依赖网络地址。":
            "当前运行拓扑为“浏览器—Next.js—FastAPI/LangGraph—本地持久化—可选 HPC”。Next.js 调用 FastAPI；任务、Checkpoint、报告和计算结果保存在项目数据目录；启用 HPC 配置时，POTCAR、SSH 私钥和 known_hosts 通过只读挂载进入后端服务。",
        "容器化实现应先在离线模式完成可重复构建，再逐步接入学术接口和集群。验收顺序包括：锁定 Python 与 Node 依赖、生成后端和前端多阶段镜像、建立非 root 用户、验证持久卷与 Secret 权限、执行 docker compose config、运行健康检查与 596 项测试函数，最后在人工创建的测试任务中验证只读预检、上传、提交、监控和下载。任何阶段未通过时，远程写入和提交开关保持关闭。":
            "Compose 默认模式通过 backend:8000 与 frontend:3000 组织服务，并为 FastAPI 配置健康检查和关键数据卷；HPC 叠加模式只读挂载密钥、known_hosts 与 PBE。远程预检、写入和提交分别由独立环境开关控制，默认配置不产生远程副作用。",
        "网页端仅开放第一批低风险人工门，外部结构上传、远程上传、sbatch、结果下载和重算尚未形成完整的网页授权闭环。":
            "前端覆盖任务历史、恢复、文件、结构、报告、DFT 监控以及远程上传和 sbatch 审查。真实远程操作由服务端环境开关与独立确认短语共同约束，连接状态探针本身不会触发上传或提交。",
        "后续迭代应遵循科学能力与工程风险同步提升的顺序。近期工作首先完成文档—代码一致性、结构文件上传和低风险网页审查闭环；随后建设作业中心、结果下载、失败诊断和高风险确认；生产化阶段再引入 PostgreSQL、对象存储、任务队列、身份认证和审计监控；科研增强阶段重点补充模型适用域评估、结构系综、竞争相与凸包、自由能修正、多中间体反应路径以及与实验数据的闭环验证。每项新增能力均需明确其证据类型、误差来源和可被否证的验收指标。":
            "现有实现通过确定性科学判据、持久化状态、人工审查和远程操作开关共同限定系统行为。文献质量、CGCNN 适用范围、DFT 收敛性和参考能一致性均在对应阶段保留来源、状态或警告，避免把流程完成等同于科学结论成立。",
        "日志采用结构化格式，并以 task_id、job_id 和 slurm_job_id 关联 API、图节点、文献检索、文件操作和集群操作。核心指标包括任务成功率、节点耗时、待审批任务数、LLM 调用错误与成本、文献接口限流次数、Slurm 状态分布、VASP 收敛率、磁盘使用量和恢复失败次数。":
            "任务状态以 task_id 为主索引，并通过 job_id 与 slurm_job_id 关联图节点、输入文件、集群作业、下载结果和恢复位置。前端时间线、作业面板和报告接口从持久化记录读取状态，使长周期任务能够在页面刷新后继续定位。",
        "持久化方案应随研究规模和并发需求分阶段演进。原型阶段可使用SQLite、JSON和受控文件目录保存任务、审批、结构与计算元数据；进入多人协作或持续运行阶段后，应将任务状态、审批记录、作业索引和版本信息迁移至PostgreSQL等事务数据库，将POSCAR、VASP输出和日志保留在具有访问控制与备份策略的文件系统或对象存储中。数据库记录需包含schema_version、created_at、updated_at和来源哈希，并通过迁移脚本维持版本兼容。":
            "当前持久化由 SQLite Checkpoint、JSON 业务记录和受控文件目录共同承担：Checkpoint 保存 LangGraph 恢复位置，JSON 记录 task_id、人工决定、job_id 与 slurm_job_id，文件目录保存 POSCAR、VASP 输入输出、报告和下载结果。记录包含 schema_version、时间戳与来源身份，用于跨进程恢复和结果追溯。",
        "日志采用结构化格式，并以 task_id、job_id 和 slurm_job_id 关联 API、图节点、文献检索、文件操作和集群操作。核心指标包括任务成功率、节点耗时、待审批任务数、LLM 调用错误与成本、文献接口限流次数、Slurm 状态分布、VASP 收敛率、磁盘使用量和恢复失败次数。":
            "任务状态以 task_id 为主索引，并通过 job_id 与 slurm_job_id 关联图节点、输入文件、集群作业、下载结果和恢复位置。前端时间线、作业面板和报告接口从持久化记录读取状态，使长周期任务能够在页面刷新后继续定位。",
    }
    for old, new in replacements.items():
        replace_paragraph(document, old, new)

    paragraph_by_text(document, "1.4 设计原则").text = "1.5 设计原则"
    paragraph_by_text(document, "1.5 相关科研智能体的方法学启示").text = "1.6 相关科研智能体的方法学启示"
    audience = paragraph_by_text(
        document,
        "系统面向电化学、催化、高熵合金和材料计算研究者，也允许缺少 VASP、Linux 或超算经验的使用者从自然语言和结构文件开始。领域熟悉者可以直接指定元素组合、POSCAR/CIF、执行范围和参数修订；计算经验较少的用户则通过术语解释、推荐选项、阶段摘要和人工确认理解每一步的目的与风险。系统降低的是工具编排和重复操作门槛，并不降低形成科研结论所需的证据标准。",
    )
    insert_sequence(audience, [
        ("Heading 2", "1.4 计算工作流门槛与设计价值"),
        ("Normal", "传统 DFT 工作流要求使用者同时理解 bulk、slab、周期性边界、吸附位点与参考态，还要正确组织 POSCAR、INCAR、KPOINTS、POTCAR 和 Slurm 脚本。元素与赝势顺序、截断能、k 点、磁性、收敛阈值、真空层和固定原子稍有不一致，就可能产生无法比较的能量或浪费超算资源。"),
        ("Normal", "进入超算后，用户还需处理 SSH 密钥、Linux 路径、sbatch、squeue/sacct、日志诊断、结果下载和断点恢复。Catalyst Agent 将这些繁琐动作转化为自然语言入口、结构化参数、可查看文件、状态解释和人工确认卡，使使用者能看懂系统将要做什么、哪些参数可改以及何时产生真实成本，同时让科学规则和外部操作继续由确定性程序约束。"),
    ])

    high_entropy = paragraph_by_text(
        document,
        "从催化设计角度看，高熵合金的优势在于组成空间和局域位点空间非常大，可通过元素组合与比例同时调节电子结构、几何应变、表面偏析和中间体吸附。然而，所谓“高熵效应”不能直接等同于高活性；真实活性还取决于材料是否形成目标固溶体、表面在电化学条件下是否重构、活性位点是否暴露以及反应路径中的吸附自由能是否适中。",
    )
    insert_after(
        high_entropy,
        "在绿色氢能、燃料电池、金属空气电池以及 CO₂/含氮小分子电化学转化中，催化剂需要同时协调多个中间体的形成、转化与脱附。高熵体系提供的大量异质局域位点使其有机会突破单一金属位点的吸附关系限制，并通过非贵金属协同改善成本与资源约束；与此同时，巨大的组合空间也使人工经验筛选迅速失效。因而，高熵电催化既是具有重要潜力的材料方向，也是适合采用文献证据、机器学习预筛和 DFT 分层验证的典型复杂问题。",
    )

    c_stage = paragraph_by_text(
        document,
        "C 阶段把候选设计、结构生成和计算验证组织为递进的资源分配过程。C1-C4 完成候选约束、评分、五元组合生成与人工选择；C5 构造 FCC bulk；C6 在生产 CGCNN、临时训练模型和 Bulk DFT 路径之间形成可审查的形成能来源；C7 联合形成能、原子尺寸差 δ 和 Ω 参数筛选可升级结构；C8-C10 生成并审查（111）slab 与 VASP 五文件；C11 执行预检、上传、SHA-256 校验、Slurm 提交、监控、下载和结果解析；C12.1-C12.7 从弛豫后的 clean slab 选择单一中间体与位点，建立吸附结构、准备并执行吸附 DFT，最后核对三项能量及参考来源。每次阶段升级均提高计算成本和结论具体性，因此需要更严格的身份校验与人工确认。",
    )
    insert_sequence(c_stage, [
        ("Heading 3", "3.4.1 五元与近等原子比建模依据"),
        ("Normal", "C1-C3 将候选固定为五元，是因为五种主元是高熵合金最常见、也最具代表性的研究起点：它满足经典多主元高熵材料的常用定义，同时避免元素数继续增加后组合空间、局域位点类型和 DFT 成本过快膨胀。对电催化而言，五元体系已经能够提供丰富的异质近邻与吸附位点，又便于从论文中追溯完整组成并进行候选间比较。因此这里的“五元”是当前工作流的明确建模边界，不表示四元或六元材料没有科学价值。"),
        ("Normal", "组成采用近等原子比，是为了保持多个元素均为主元并提高理想混合构型熵，避免某一元素占比过低而退化为微量掺杂模型。当前 FCC 2x2x2 超胞含 32 个原子，不能被 5 整除，因此无特殊元素时使用 7/7/6/6/6，并轮换两个 7 原子元素；含 Cu 时使用 Cu8 与其余元素各 6 的 8/6/6/6/6 规则。两者都围绕 20 at.% 的理想等原子组成作最小整数近似，而不是宣称已经得到实验最优比例。"),
        ("Heading 3", "3.4.2 P 区元素限额与 FCC 起始假设"),
        ("Normal", "当前模型支持集合中的 P 区元素按代码定义为 Al、Zn、Ga、Ge，单个候选最多包含其中一种。其原因是本工作流以金属过渡元素主导的置换型 FCC 固溶体为起始模型；同时引入多种 P 区元素会增加原子尺寸、价电子特征、键合方式和偏析倾向的差异，可能提高形成复杂金属间化合物、第二相或非理想局域结构的风险，从而削弱简单随机固溶体假设。为降低其结构扰动，含 Cu 与 P 区元素时使用 Cu8、P 区元素 3、其余三种各 7；不含 Cu 时使用 P 区元素 4、其余四种各 7。"),
        ("Normal", "FCC 被用作过渡金属高熵合金常见且便于比较的初始晶格，Vegard 规则用于给出起始晶格常数。P 区限额、近等比例和 FCC 均是缩小搜索空间的建模先验；随后仍需通过 C6 形成能、C7 的 δ/Ω 判据、bulk/slab DFT 及实验表征验证，不能把这些先验直接解释为单相形成或催化活性的证明。"),
    ])

    paragraph_by_text(document, "3.5 节点设计规范").text = "3.5 节点设计规范"
    add_tool_section(document)
    add_workflow_completeness_section(document)

    for text in (
        "容器以非 root 用户运行，不将 .env、私钥、POTCAR、运行结果写入镜像；",
        "镜像版本与应用版本对应，生产部署禁止使用不可追踪的 latest 标签。",
        "上传文件限制扩展名、大小和目标目录，使用服务端生成文件名防止路径穿越；",
        "API 实施登录、任务归属、角色权限、速率限制和审计日志；",
        "SQLite Checkpoint、JSON 业务记录和单线程 TaskManager 适合研发验证，不适合高并发、多租户和跨节点故障恢复。",
    ):
        remove_paragraph(document, text)

    paragraph_by_text(document, "10 系统限制与迭代计划").text = "10 科学边界与已实现能力"
    paragraph_by_text(document, "10.2 迭代计划").text = "10.2 已实现能力汇总"
    paragraph_by_text(document, "10.3 初稿待确认事项").text = "10.3 实现依据"
    for text in (
        "项目正式名称、单位、负责人、成员和文档署名；",
        "生产服务器区域、域名、认证方案和预期用户规模；",
        "FastAPI 已实现接口与本设计建议接口的最终对应关系；",
        "Checkpoint、JSON、SQLite 或 PostgreSQL 的最终持久化方案；",
        "CGCNN 模型版本、适用元素域和预测误差的正式记录方式；",
        "集群账号、分区、VASP 启动命令和 POTCAR 管理制度。",
    ):
        remove_paragraph(document, text)
    insert_after(
        paragraph_by_text(document, "10.3 实现依据"),
        "本说明书以当前 README、FastAPI 路由、LangGraph 节点、领域服务、前端组件、Docker 配置、测试目录和已持久化任务记录为实现依据。文档中的功能描述均可回溯到上述代码或运行产物。",
    )

    production_switch = paragraph_by_text(document, "8.6 生产环境开关")
    production_switch.text = "8.6 远程操作开关"
    switch_body = insert_after(
        production_switch,
        "网页远程操作只有在 WEB_REMOTE_OPERATIONS_ENABLED、CLUSTER_REMOTE_WRITE_ENABLED 和 CLUSTER_SUBMISSION_ENABLED 等安全条件同时满足时才可执行。连接状态的绿色指示只表示最小 Kimi 请求或只读 SSH 探针成功，不代表用户已批准任何上传或提交。",
    )


def update_tables(document: Document, version: str) -> None:
    document.tables[0].cell(5, 1).text = version
    set_table_data(table_by_headers(document, ["方法与路径", "用途", "关键约束"]), [
        ["方法与路径", "用途", "当前状态与约束"],
        ["POST /api/tasks", "创建自然语言科研任务", "已实现；202 Accepted，后台执行并立即返回 task_id"],
        ["GET /api/tasks/{task_id}", "状态、时间线、审查和恢复信息", "已实现；聊天台和右侧工作台共同使用"],
        ["POST /api/tasks/{task_id}/reviews", "提交 B6-C12.7 人工决定", "已实现；校验 review_id、类型、对象、幂等键和高风险确认短语"],
        ["GET /api/tasks/{task_id}/files|structures", "文件、结构预览与下载", "已实现；POTCAR 禁止下载，结构使用网页三维查看"],
        ["GET /api/tasks/{task_id}/jobs", "Slurm/VASP 状态与日志", "已实现只读查询、刷新和日志接口"],
        ["POST|GET /api/tasks/{task_id}/report", "生成并下载任务报告", "已实现 HTML、Markdown 与 JSON 报告"],
        ["GET|POST /api/system/connections", "Kimi/HPC 配置与实时探针", "已实现；返回脱敏状态，不产生远程写入"],
    ])
    set_table_data(table_by_headers(document, ["容器", "职责", "是否持有敏感信息"]), [
        ["部署单元", "职责", "当前状态与敏感信息"],
        ["backend", "FastAPI、LangGraph、文献、结构、CGCNN 和集群服务", "Dockerfile 已提供；运行时读取 LLM Key，HPC 模式只读挂载 SSH/PBE"],
        ["frontend", "Next.js/assistant-ui、聊天台与工作台", "多阶段 Dockerfile 已提供；不保存密钥和 POTCAR"],
        ["docker-compose.yml", "默认安全演示", "已提供；强制关闭网页远程写入和提交"],
        ["docker-compose.hpc.yml", "管理员真实计算叠加配置", "已提供；只读挂载密钥、known_hosts 和许可 PBE"],
    ])
    set_table_data(table_by_headers(document, ["挂载对象", "容器路径示例", "模式"]), [
        ["挂载对象", "容器路径", "模式"],
        ["Checkpoint", "/app/data/checkpoints", "持久化读写卷"],
        ["工作流运行记录", "/app/data/workflow_runs", "持久化读写卷"],
        ["报告与集群记录", "/app/data/reports、cluster_jobs、cluster_results", "持久化读写卷"],
        ["POTCAR/PBE", "/app/database/PBE", "只读，受许可证约束"],
        ["SSH 私钥", "/run/secrets/cluster_key", "只读"],
        ["known_hosts", "/run/secrets/known_hosts", "只读"],
    ])
    set_table_data(table_by_headers(document, ["配置项", "默认值", "开启条件"]), [
        ["配置项", "默认值", "开启条件"],
        ["WEB_REMOTE_OPERATIONS_ENABLED", "false", "管理员可信环境；网页端仍保留独立人工门"],
        ["CLUSTER_PREFLIGHT_ENABLED", "false", "SSH、known_hosts 与远程目录配置经验证"],
        ["CLUSTER_REMOTE_WRITE_ENABLED", "false", "上传对象、远程路径与 SHA-256 已审查"],
        ["CLUSTER_SUBMISSION_ENABLED", "false", "上传校验通过并输入精确 SUBMIT 短语"],
        ["LLM_ENABLED", "按环境", "Kimi Key、模型、超时与规则回退配置完成"],
    ])
    set_table_data(table_by_headers(document, ["阶段", "建设内容", "验收重点"]), [
        ["已实现层级", "现有内容", "设计作用"],
        ["当前完整工作台", "A1-C12.7 时间线、审查卡、历史恢复、文件、结构、报告和 DFT 监控", "三个入口可进入同一身份链，工作流内容保留在聊天台"],
        ["连接与远程安全", "Kimi/HPC 实时状态、网页上传和 sbatch 人工门", "只读探针与真实副作用分离，确认短语不可绕过"],
        ["运行与容器配置", "Python/Next.js 本地运行、后端/前端 Dockerfile、默认和 HPC Compose", "统一前后端入口，并隔离许可文件与远程凭据"],
        ["科研可追溯性", "文献来源、模型来源、结构身份、输入摘要、Slurm 标识和能量来源", "让每个结论可定位到对应 task_id 与计算对象"],
        ["结果交付", "文件查看、三维结构、状态日志、下载、吸附能审查和任务报告", "把分散的科研产物组织为可检查的任务记录"],
    ])
    set_table_data(table_by_headers(document, ["配置类别", "示例配置", "说明"]), [
        ["配置类别", "示例配置", "说明"],
        ["Kimi", "LLM_ENABLED、LLM_API_KEY、LLM_BASE_URL、LLM_MODEL", "仅保存在服务端环境或 Secret；Kimi K3 使用 temperature=1"],
        ["文献服务", "CROSSREF_MAILTO、SEMANTIC_SCHOLAR_API_KEY", "支持限流、来源追踪和本地优先召回"],
        ["集群", "CLUSTER_SSH_HOST、PORT、USER、REMOTE_RUNS_ROOT", "不得由普通用户自由修改"],
        ["网页安全", "WEB_REMOTE_OPERATIONS_ENABLED、REMOTE_WRITE、SUBMISSION", "默认关闭，管理员显式开启，人工门继续生效"],
        ["路径与许可", "DATA_ROOT、PBE_ROOT、RESULT_ROOT", "绝对路径边界校验；PBE 不写入镜像"],
    ])


def format_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    for style_name, east_asia, size, bold, color, before, after in (
        ("Normal", BODY_EAST_ASIA, 10.5, False, RGBColor(0, 0, 0), 0, 6),
        ("Heading 1", HEADING_EAST_ASIA, 16, True, ACCENT, 16, 8),
        ("Heading 2", HEADING_EAST_ASIA, 14, True, INK, 12, 6),
        ("Heading 3", HEADING_EAST_ASIA, 12, True, INK, 9, 4),
    ):
        style = document.styles[style_name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = style_name.startswith("Heading")
        if style_name == "Normal":
            style.paragraph_format.line_spacing = 1.25

    for paragraph in document.paragraphs:
        heading = paragraph.style.name.startswith("Heading")
        toc = paragraph.style.name.startswith("toc")
        east_asia = HEADING_EAST_ASIA if heading else BODY_EAST_ASIA
        for run in paragraph.runs:
            set_run_fonts(run, east_asia)
            if heading:
                run.bold = True
            if toc:
                run.font.size = Pt(10)
        if paragraph.style.name == "Normal" and paragraph.text.strip():
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = Pt(6)
        if heading:
            paragraph.paragraph_format.keep_with_next = True

    for table in document.tables:
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = tc_pr.first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                for side, value in (("top", 80), ("bottom", 80), ("start", 100), ("end", 100)):
                    node = tc_mar.find(qn(f"w:{side}"))
                    if node is None:
                        node = OxmlElement(f"w:{side}")
                        tc_mar.append(node)
                    node.set(qn("w:w"), str(value))
                    node.set(qn("w:type"), "dxa")
                if row_index == 0:
                    shading = tc_pr.find(qn("w:shd"))
                    if shading is None:
                        shading = OxmlElement("w:shd")
                        tc_pr.append(shading)
                    shading.set(qn("w:fill"), "E8EEF5")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.12
                    for run in paragraph.runs:
                        set_run_fonts(run, BODY_EAST_ASIA)
                        run.font.size = Pt(9)
                        if row_index == 0:
                            run.bold = True

    for paragraph in document.sections[0].header.paragraphs + document.sections[0].footer.paragraphs:
        for run in paragraph.runs:
            set_run_fonts(run, BODY_EAST_ASIA)


def format_cover(document: Document, showcase: bool) -> None:
    paragraphs = document.paragraphs
    paragraphs[1].text = "AI4S 科研智能体项目"
    paragraphs[2].text = "智能体设计方案与设计说明书"
    paragraphs[3].text = (
        "面向电催化高熵合金设计与计算的一体化智能工作台"
        if showcase
        else "面向电催化高熵合金研究的可追溯科研工作流智能体"
    )
    paragraphs[4].text = "Catalyst Agent"
    paragraphs[7].text = (
        "以自然语言连接文献证据、候选设计、结构建模、VASP、超算和结果报告，使复杂计算步骤可理解、可确认、可恢复。"
        if showcase
        else "本文件描述系统目标、科学依据、已实现工作流、数据契约、人机协同和安全边界。"
    )
    specs = {
        1: (12, False, MUTED),
        2: (22, True, INK),
        3: (15, True, ACCENT),
        4: (18, True, INK),
        7: (11, False, MUTED),
    }
    for index, (size, bold, color) in specs.items():
        paragraph = paragraphs[index]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_fonts(run, HEADING_EAST_ASIA if bold else BODY_EAST_ASIA)
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color
    document.tables[0].cell(5, 1).text = (
        "V1.0 项目展示版（已实现功能修订稿）"
        if showcase
        else "V1.0 专业设计版（已实现功能修订稿）"
    )


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build(source: Path, professional: Path, showcase: Path) -> None:
    base = Document(source)
    apply_content_updates(base)
    update_tables(base, "V1.0 已实现功能修订稿")
    format_document(base)
    set_update_fields(base)

    professional_doc = copy.deepcopy(base)
    format_cover(professional_doc, showcase=False)
    professional_doc.save(professional)

    showcase_doc = Document(professional)
    format_cover(showcase_doc, showcase=True)
    showcase_doc.save(showcase)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("professional", type=Path)
    parser.add_argument("showcase", type=Path)
    args = parser.parse_args()
    args.professional.parent.mkdir(parents=True, exist_ok=True)
    build(args.source, args.professional, args.showcase)
    print(args.professional)
    print(args.showcase)


if __name__ == "__main__":
    main()
